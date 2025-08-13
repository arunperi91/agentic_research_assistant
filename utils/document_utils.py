import os
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from workflow.state import ResearchState

def create_professional_word_doc(report_content: str, topic: str, state: ResearchState = None) -> str:
    """Create a professional Word document from report content"""
    
    # Create document
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'Research Report: {topic}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    if state:
        metadata_para = doc.add_paragraph()
        metadata_para.add_run('Generated: ').bold = True
        metadata_para.add_run(datetime.now().strftime("%B %d, %Y"))
        
        if state.get('quality_score'):
            metadata_para.add_run('\nQuality Score: ').bold = True
            metadata_para.add_run(f"{state['quality_score']:.2f}/1.00")
        
        if state.get('processing_time'):
            metadata_para.add_run('\nProcessing Time: ').bold = True
            metadata_para.add_run(f"{state['processing_time']:.1f} seconds")
    
    doc.add_paragraph()  # Add space
    
    # Process report content
    sections = report_content.split('\n# ')
    
    for i, section in enumerate(sections):
        if not section.strip():
            continue
            
        lines = section.strip().split('\n')
        if not lines:
            continue
            
        # First section doesn't need heading processing if it starts with title
        if i == 0 and lines[0].startswith('# Research Report:'):
            lines = lines[1:]  # Skip title as we already added it
        
        # Process section heading
        section_title = lines[0] if lines else ""
        if section_title:
            if section_title.startswith('# '):
                section_title = section_title[2:]
            doc.add_heading(section_title, 1)
        
        # Process section content
        current_para = None
        for line in lines[1:]:
            line = line.strip()
            
            if not line:
                if current_para:
                    current_para = None
                continue
            
            # Handle subheadings
            if line.startswith('## '):
                doc.add_heading(line[3:], 2)
                current_para = None
            elif line.startswith('### '):
                doc.add_heading(line[4:], 3)
                current_para = None
            # Handle bullet points
            elif line.startswith('- ') or line.startswith('• '):
                bullet_para = doc.add_paragraph(line[2:], style='List Bullet')
                current_para = None
            # Handle numbered lists
            elif line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
                num_para = doc.add_paragraph(line[3:], style='List Number')
                current_para = None
            # Handle bold text markers
            elif line.startswith('**') and line.endswith('**'):
                bold_para = doc.add_paragraph()
                bold_para.add_run(line[2:-2]).bold = True
                current_para = None
            # Regular text
            else:
                if current_para is None:
                    current_para = doc.add_paragraph()
                
                if current_para.text:
                    current_para.add_run(' ')
                current_para.add_run(line)
        
        # Add space between sections
        doc.add_paragraph()
    
    # Save to temporary file
    temp_dir = tempfile.gettempdir()
    filename = f"research_report_{topic.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(temp_dir, filename)
    
    doc.save(filepath)
    return filepath
